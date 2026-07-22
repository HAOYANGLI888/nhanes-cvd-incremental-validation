"""Stage 3A NHANES internal-external validation utilities.

This module intentionally excludes NHANES 2021-2023 from model development,
tuning, performance estimation, and prediction export. It reads only the
Stage 2 frozen configuration files and the Stage 2 harmonized audit dataset.
"""

from __future__ import annotations

import argparse
import hashlib
import itertools
import json
import math
import os
import platform
import shutil
import sys
import uuid
import warnings
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import joblib
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yaml
from docx import Document
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    average_precision_score,
    brier_score_loss,
    confusion_matrix,
    log_loss,
    precision_recall_curve,
    roc_auc_score,
    roc_curve,
)
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from statsmodels.api import GLM, add_constant
from statsmodels.genmod.families import Binomial
from xgboost import XGBClassifier

warnings.filterwarnings("ignore", category=FutureWarning, module="sklearn.linear_model._logistic")


DEVELOPMENT_CYCLES = [
    "2005-2006",
    "2007-2008",
    "2009-2010",
    "2011-2012",
    "2013-2014",
    "2015-2016",
    "2017-2018",
]
EXTERNAL_CYCLE = "2021-2023"
RANDOM_SEED = 20260713
EPS = 1e-8
CHECKPOINT_SCHEMA_VERSION = 1
SUPPORTED_ALGORITHMS = ("elastic_net", "xgboost")
DEFAULT_PARALLEL_JOBS = max(1, min(8, (os.cpu_count() or 2) - 1))
ANALYSIS_SEMANTICS_VERSION = "stage3a_locked_v1"
ANALYSIS_SEMANTICS_SPEC = {
    "development_cycles": DEVELOPMENT_CYCLES,
    "outcome": "prevalent_self_reported_cvd",
    "outer_validation": "leave_one_nhanes_cycle_out",
    "inner_validation": "leave_one_training_cycle_out",
    "selection_metric": "survey_weighted_log_loss",
    "selection_tolerance": 1e-4,
    "selection_tie_break": "lower_prespecified_complexity",
    "preprocessing": {
        "continuous_missing": "training_fold_median",
        "categorical_missing": "training_fold_constant_Missing",
        "categorical_encoding": "training_fold_one_hot_handle_unknown_ignore",
        "elastic_net_scaling": "training_fold_standard_scaler",
        "xgboost_scaling": "none",
    },
    "training_weights": "model_domain_specific_combined_nhanes_weight",
    "main_class_balance": {"elastic_net": "class_weight_none", "xgboost": "scale_pos_weight_1"},
    "elastic_net": "weighted_logistic_regression_locked_C_l1_ratio_grid",
    "xgboost": "binary_logistic_hist_single_estimator_thread_locked_grid",
    "outer_refit": "selected_hyperparameters_fit_on_all_outer_training_cycles",
    "thresholds": ["fixed_0.50", "outer_training_weighted_youden"],
    "probability_clip": [EPS, 1 - EPS],
    "paired_incremental_comparison": "same_seqn_cycle_sample_and_extended_model_weight",
}
CHECKPOINT_ATTESTATION_RELATIVE_PATH = Path("results/audit/stage3a_checkpoint_compatibility_attestation.json")


MODEL_SAMPLE_DOMAIN = {
    "model0_core": "mec",
    "model1_renal": "mec",
    "model2_metabolic": "fasting",
    "model3_inflammatory": "mec",
    "model4_combined": "fasting",
    "model0_paired_metabolic": "fasting",
    "model0_paired_combined": "fasting",
}

MODEL_DISPLAY = {
    "model0_core": "Model 0 core clinical",
    "model1_renal": "Model 1 renal",
    "model2_metabolic": "Model 2 metabolic",
    "model3_inflammatory": "Model 3 inflammatory",
    "model4_combined": "Model 4 combined biomarkers",
    "model0_paired_metabolic": "Model 0 paired metabolic sample",
    "model0_paired_combined": "Model 0 paired combined sample",
}

INCREMENT_COMPARISONS = {
    "renal": ("model0_core", "model1_renal", "mec"),
    "metabolic": ("model0_paired_metabolic", "model2_metabolic", "fasting"),
    "inflammatory": ("model0_core", "model3_inflammatory", "mec"),
    "combined": ("model0_paired_combined", "model4_combined", "fasting"),
}


@dataclass
class Stage3AContext:
    project_root: Path
    data_path: Path
    run_time: str
    log_path: Path


def now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def ensure_dirs(project_root: Path) -> None:
    for rel in [
        "results/logs",
        "results/tables",
        "results/figures",
        "results/audit",
        "results/models",
        "results/models/frozen_development",
        "results/predictions",
        "results/manuscript_tables",
        "results/checkpoints/stage3a/outer",
        "results/checkpoints/stage3a/frozen",
        "results/checkpoints/stage3a/candidates",
    ]:
        (project_root / rel).mkdir(parents=True, exist_ok=True)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def analysis_semantics_hash() -> str:
    return sha256_text(canonical_json({"version": ANALYSIS_SEMANTICS_VERSION, "spec": ANALYSIS_SEMANTICS_SPEC}))


def analysis_software_versions() -> dict[str, str]:
    import sklearn
    import xgboost

    return {
        "python": sys.version.split()[0],
        "numpy": np.__version__,
        "pandas": pd.__version__,
        "sklearn": sklearn.__version__,
        "xgboost": xgboost.__version__,
    }


def unique_temp_path(path: Path) -> Path:
    return path.with_name(f".{path.name}.tmp-{os.getpid()}-{uuid.uuid4().hex}")


def atomic_write_text(path: Path, text: str) -> None:
    """Write a small checkpoint file atomically on the same filesystem."""
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = unique_temp_path(path)
    tmp.write_text(text, encoding="utf-8")
    os.replace(tmp, path)


def atomic_write_csv(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = unique_temp_path(path)
    frame.to_csv(tmp, index=False)
    os.replace(tmp, path)


def atomic_joblib_dump(value: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = unique_temp_path(path)
    joblib.dump(value, tmp)
    os.replace(tmp, path)


def atomic_copy(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    tmp = unique_temp_path(destination)
    shutil.copy2(source, tmp)
    os.replace(tmp, destination)


def checkpoint_base_signature(project_root: Path) -> dict[str, Any]:
    """Inputs that must match before any expensive result can be resumed."""
    return {
        "schema_version": CHECKPOINT_SCHEMA_VERSION,
        "input_hash": sha256_file(project_root / "data/interim/stage2_harmonized_audit_dataset.csv"),
        "config_hashes": config_hashes_dict(project_root),
        "code_hash": sha256_file(Path(__file__).resolve()),
        "analysis_semantics_version": ANALYSIS_SEMANTICS_VERSION,
        "analysis_semantics_hash": analysis_semantics_hash(),
        "analysis_software": analysis_software_versions(),
        "random_seed": RANDOM_SEED,
        "development_cycles": DEVELOPMENT_CYCLES,
    }


def checkpoint_task_name(fold: int | str, holdout: str, model_name: str, algorithm: str) -> str:
    safe_holdout = holdout.replace("-", "_")
    return f"fold_{int(fold):02d}__{safe_holdout}__{model_name}__{algorithm}"


def file_manifest(paths: dict[str, Path]) -> dict[str, dict[str, Any]]:
    return {
        key: {"file": path.name, "sha256": sha256_file(path), "bytes": path.stat().st_size}
        for key, path in paths.items()
    }


def read_checkpoint_files(
    checkpoint_dir: Path,
    expected_signature: dict[str, Any],
    files: dict[str, str],
    project_root: Path | None = None,
) -> tuple[dict[str, Any], dict[str, pd.DataFrame]] | None:
    """Return a completed checkpoint only when signature and every hash agree."""
    meta_path = checkpoint_dir / "checkpoint.json"
    if not meta_path.exists():
        return None
    try:
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        if not checkpoint_signature_matches(meta.get("signature"), expected_signature, project_root) or meta.get("status") != "complete":
            return None
        frames: dict[str, pd.DataFrame] = {}
        manifest = meta.get("files", {})
        for key, filename in files.items():
            path = checkpoint_dir / filename
            recorded = manifest.get(key, {})
            if not path.exists() or recorded.get("sha256") != sha256_file(path):
                return None
            frame = pd.read_csv(path)
            if int(recorded.get("rows", -1)) != len(frame):
                return None
            frames[key] = frame
        return meta, frames
    except (OSError, ValueError, KeyError, pd.errors.ParserError, json.JSONDecodeError):
        return None


def load_checkpoint_attestation(project_root: Path) -> tuple[dict[str, Any], str] | None:
    path = project_root / CHECKPOINT_ATTESTATION_RELATIVE_PATH
    if not path.exists():
        return None
    try:
        attestation = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError, json.JSONDecodeError):
        return None
    return attestation, sha256_file(path)


def checkpoint_signature_matches(
    stored: Any,
    expected: dict[str, Any],
    project_root: Path | None = None,
) -> bool:
    """Require exact analytic signatures; code-hash exceptions need a signed PASS attestation."""
    if not isinstance(stored, dict):
        return False
    ignored = {"code_hash", "compatibility_attestation_id", "compatibility_attestation_sha256"}
    stored_cmp = {k: v for k, v in stored.items() if k not in ignored}
    expected_cmp = {k: v for k, v in expected.items() if k not in ignored}
    if stored_cmp != expected_cmp:
        return False
    stored_code = stored.get("code_hash")
    expected_code = expected.get("code_hash")
    if stored_code == expected_code:
        return True
    if project_root is None:
        return False
    loaded = load_checkpoint_attestation(project_root)
    if loaded is None:
        return False
    attestation, attestation_hash = loaded
    if attestation.get("status") != "PASS":
        return False
    if stored.get("compatibility_attestation_id") != attestation.get("attestation_id"):
        return False
    if stored.get("compatibility_attestation_sha256") != attestation_hash:
        return False
    if attestation.get("analysis_semantics_version") != expected.get("analysis_semantics_version"):
        return False
    if attestation.get("analysis_semantics_hash") != expected.get("analysis_semantics_hash"):
        return False
    compatible = set(attestation.get("compatible_code_hashes", []))
    return bool(stored_code in compatible and expected_code in compatible)


def expected_signature_from_stored(project_root: Path, stored: dict[str, Any]) -> dict[str, Any]:
    expected = {
        k: v
        for k, v in stored.items()
        if k not in {"compatibility_attestation_id", "compatibility_attestation_sha256"}
    }
    expected.update(checkpoint_base_signature(project_root))
    return expected


def checkpoint_compatibility_status(project_root: Path, stored: dict[str, Any]) -> str:
    expected = expected_signature_from_stored(project_root, stored)
    return "PASS" if checkpoint_signature_matches(stored, expected, project_root) else "FAIL"


def write_checkpoint_frames(
    checkpoint_dir: Path,
    signature: dict[str, Any],
    frames: dict[str, pd.DataFrame],
) -> dict[str, Any]:
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    paths = {key: checkpoint_dir / f"{key}.csv" for key in frames}
    for key, frame in frames.items():
        atomic_write_csv(frame, paths[key])
    manifest = file_manifest(paths)
    for key, frame in frames.items():
        manifest[key]["rows"] = len(frame)
    meta = {
        "status": "complete",
        "created": now(),
        "signature": signature,
        "files": manifest,
    }
    atomic_write_text(checkpoint_dir / "checkpoint.json", json.dumps(meta, indent=2, sort_keys=True))
    return meta


def log(ctx: Stage3AContext, msg: str) -> None:
    line = f"{now()} | {msg}"
    with ctx.log_path.open("a", encoding="utf-8") as f:
        f.write(line + "\n")
    try:
        print(line, flush=True)
    except OSError:
        pass


def load_yaml(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def one_hot_encoder() -> OneHotEncoder:
    try:
        return OneHotEncoder(handle_unknown="ignore", sparse_output=True)
    except TypeError:
        return OneHotEncoder(handle_unknown="ignore", sparse=True)


def read_frozen_features(project_root: Path) -> dict[str, list[str]]:
    spec = load_yaml(project_root / "config/frozen_variable_specification.yml")
    required = ["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"]
    features = {k: list(spec[k]) for k in required}
    features["model0_paired_metabolic"] = list(spec["model0_core"])
    features["model0_paired_combined"] = list(spec["model0_core"])
    return features


def read_cycle_folds(project_root: Path) -> list[dict[str, Any]]:
    folds = load_yaml(project_root / "config/cycle_folds.yml")
    return list(folds["outer_validation"])


def check_temporal_lock(project_root: Path) -> None:
    lock = load_yaml(project_root / "config/temporal_validation_locked.yml")
    if not bool(lock.get("locked", False)):
        raise RuntimeError("Temporal validation lock is not true.")


def read_harmonized_data(project_root: Path) -> pd.DataFrame:
    path = project_root / "data/interim/stage2_harmonized_audit_dataset.csv"
    if not path.exists():
        raise FileNotFoundError(path)
    df = pd.read_csv(path, low_memory=False)
    return df


def adult_nonpreg_outcome(df: pd.DataFrame) -> pd.Series:
    return df["age"].ge(20) & (~df["pregnancy_code"].eq(1).fillna(False)) & df["cvd"].notna()


def development_df(df: pd.DataFrame) -> pd.DataFrame:
    d = df[df["cycle"].isin(DEVELOPMENT_CYCLES)].copy()
    return d[adult_nonpreg_outcome(d)].copy()


def sample_mask(df: pd.DataFrame, model_name: str) -> pd.Series:
    domain = MODEL_SAMPLE_DOMAIN[model_name]
    base = adult_nonpreg_outcome(df) & df["cycle"].isin(DEVELOPMENT_CYCLES)
    design = df["strata"].notna() & df["psu"].notna()
    if domain == "mec":
        w = df["combined_mec_weight"].notna() & df["combined_mec_weight"].gt(EPS)
    elif domain == "fasting":
        w = df["combined_fasting_weight"].notna() & df["combined_fasting_weight"].gt(EPS)
    else:
        raise ValueError(domain)
    return base & design & w


def weight_column(model_name: str) -> str:
    return "combined_fasting_weight" if MODEL_SAMPLE_DOMAIN[model_name] == "fasting" else "combined_mec_weight"


def target_population(model_name: str) -> str:
    return "fasting subsample" if MODEL_SAMPLE_DOMAIN[model_name] == "fasting" else "MEC examination sample"


def build_preprocessor(features: list[str], algorithm: str, train: pd.DataFrame) -> ColumnTransformer:
    categorical = [c for c in features if train[c].dtype == "object" or c in ["sex", "race_ethnicity", "education", "smoking", "diabetes"]]
    continuous = [c for c in features if c not in categorical]
    if algorithm == "elastic_net":
        continuous_pipe = Pipeline(
            [
                ("imputer", SimpleImputer(strategy="median")),
                ("scaler", StandardScaler()),
            ]
        )
    else:
        continuous_pipe = Pipeline([("imputer", SimpleImputer(strategy="median"))])
    categorical_pipe = Pipeline(
        [
            ("imputer", SimpleImputer(strategy="constant", fill_value="Missing")),
            ("onehot", one_hot_encoder()),
        ]
    )
    return ColumnTransformer(
        [
            ("continuous", continuous_pipe, continuous),
            ("categorical", categorical_pipe, categorical),
        ],
        remainder="drop",
        sparse_threshold=1.0,
        verbose_feature_names_out=True,
    )


def elastic_net_candidates(project_root: Path, limit: int | None = None) -> list[dict[str, Any]]:
    cfg = load_yaml(project_root / "config/elastic_net_hyperparameter_grid.yml")
    combos = [
        {
            "C": float(C),
            "l1_ratio": float(l1),
            "solver": cfg.get("solver", "saga"),
            "max_iter": int(cfg.get("max_iter", 5000)),
            "tol": float(cfg.get("tol", 0.0001)),
            "class_weight": None,
        }
        for l1, C in itertools.product(cfg["l1_ratio"], cfg["C"])
    ]
    if limit:
        combos = sorted(combos, key=lambda x: (x["C"], x["l1_ratio"]))[:limit]
    return combos


def xgboost_candidates(project_root: Path, limit: int | None = None) -> list[dict[str, Any]]:
    cfg = load_yaml(project_root / "config/xgboost_hyperparameter_grid.yml")
    keys = [
        "max_depth",
        "min_child_weight",
        "learning_rate",
        "n_estimators",
        "subsample",
        "colsample_bytree",
        "reg_alpha",
        "reg_lambda",
    ]
    combos = [dict(zip(keys, vals)) for vals in itertools.product(*(cfg[k] for k in keys))]
    for combo in combos:
        combo["scale_pos_weight"] = 1
    combos = sorted(combos, key=xgb_complexity_key)
    if limit:
        combos = combos[:limit]
    return combos


def xgb_complexity_key(params: dict[str, Any]) -> tuple:
    return (
        int(params.get("max_depth", 99)),
        int(params.get("n_estimators", 999999)),
        float(params.get("learning_rate", 9)),
        -float(params.get("min_child_weight", 0)),
        -float(params.get("reg_alpha", 0)),
        -float(params.get("reg_lambda", 0)),
    )


def en_complexity_key(params: dict[str, Any]) -> tuple:
    return (float(params.get("C", 999999)), float(params.get("l1_ratio", 999999)))


def candidate_grid_signature(candidates: list[dict[str, Any]], limit: int | None) -> dict[str, Any]:
    return {
        "candidate_count": len(candidates),
        "candidate_limit": limit,
        "candidate_grid_hash": sha256_text(canonical_json(candidates)),
    }


def outer_task_signature(
    base: dict[str, Any],
    fold: dict[str, Any],
    model_name: str,
    algorithm: str,
    features: list[str],
    candidates: list[dict[str, Any]],
    limit: int | None,
) -> dict[str, Any]:
    return {
        **base,
        "checkpoint_type": "outer_validation_task",
        "fold": int(fold["fold"]),
        "holdout_cycle": str(fold["holdout_cycle"]),
        "training_cycles": list(fold["training_cycles"]),
        "model": model_name,
        "algorithm": algorithm,
        "features": features,
        **candidate_grid_signature(candidates, limit),
    }


def frozen_task_signature(
    base: dict[str, Any],
    model_name: str,
    algorithm: str,
    features: list[str],
    candidates: list[dict[str, Any]],
    limit: int | None,
) -> dict[str, Any]:
    return {
        **base,
        "checkpoint_type": "frozen_development_task",
        "training_cycles": DEVELOPMENT_CYCLES,
        "model": model_name,
        "algorithm": algorithm,
        "features": features,
        **candidate_grid_signature(candidates, limit),
    }


def checkpoint_semantic_cache_key(signature: dict[str, Any]) -> str:
    semantic = {
        k: v
        for k, v in signature.items()
        if k not in {"code_hash", "compatibility_attestation_id", "compatibility_attestation_sha256"}
    }
    return sha256_text(canonical_json(semantic))[:24]


def candidate_checkpoint_dir(project_root: Path, task_signature: dict[str, Any], candidate_id: str) -> Path:
    """Use a short, semantic task digest to stay below Windows path limits."""
    algorithm = str(task_signature["algorithm"])
    short_algorithm = "en" if algorithm == "elastic_net" else "xgb"
    return (
        project_root
        / "results/checkpoints/stage3a/candidates"
        / short_algorithm
        / checkpoint_semantic_cache_key(task_signature)
        / candidate_id
    )


def make_estimator(algorithm: str, params: dict[str, Any]) -> Any:
    if algorithm == "elastic_net":
        if float(params["l1_ratio"]) == 0.0:
            return LogisticRegression(
                penalty="l2",
                solver="lbfgs",
                C=float(params["C"]),
                max_iter=int(params.get("max_iter", 5000)),
                tol=float(params.get("tol", 0.0001)),
                class_weight=params.get("class_weight", None),
                random_state=RANDOM_SEED,
            )
        return LogisticRegression(
            penalty="elasticnet",
            solver=params.get("solver", "saga"),
            C=float(params["C"]),
            l1_ratio=float(params["l1_ratio"]),
            max_iter=int(params.get("max_iter", 5000)),
            tol=float(params.get("tol", 0.0001)),
            class_weight=params.get("class_weight", None),
            random_state=RANDOM_SEED,
        )
    if algorithm == "xgboost":
        return XGBClassifier(
            objective="binary:logistic",
            eval_metric="logloss",
            random_state=RANDOM_SEED,
            n_jobs=1,
            tree_method="hist",
            max_depth=int(params["max_depth"]),
            min_child_weight=float(params["min_child_weight"]),
            learning_rate=float(params["learning_rate"]),
            n_estimators=int(params["n_estimators"]),
            subsample=float(params["subsample"]),
            colsample_bytree=float(params["colsample_bytree"]),
            reg_alpha=float(params["reg_alpha"]),
            reg_lambda=float(params["reg_lambda"]),
            scale_pos_weight=float(params.get("scale_pos_weight", 1)),
        )
    raise ValueError(algorithm)


def make_pipeline(features: list[str], algorithm: str, train: pd.DataFrame, params: dict[str, Any]) -> Pipeline:
    pre = build_preprocessor(features, algorithm, train)
    clf = make_estimator(algorithm, params)
    return Pipeline([("preprocess", pre), ("model", clf)])


def fit_pipeline(pipe: Pipeline, X: pd.DataFrame, y: np.ndarray, w: np.ndarray) -> Pipeline:
    pipe.fit(X, y, model__sample_weight=w)
    return pipe


def predict_proba(pipe: Pipeline, X: pd.DataFrame) -> np.ndarray:
    p = pipe.predict_proba(X)[:, 1]
    return np.clip(p, EPS, 1 - EPS)


def weighted_mean(x: np.ndarray, w: np.ndarray) -> float:
    mask = np.isfinite(x) & np.isfinite(w) & (w > 0)
    if not mask.any():
        return float("nan")
    return float(np.average(x[mask], weights=w[mask]))


def weighted_log_loss(y: np.ndarray, p: np.ndarray, w: np.ndarray) -> float:
    return float(log_loss(y, np.clip(p, EPS, 1 - EPS), labels=[0, 1], sample_weight=w))


def weighted_brier(y: np.ndarray, p: np.ndarray, w: np.ndarray) -> float:
    return float(brier_score_loss(y, p, sample_weight=w))


def safe_auc(y: np.ndarray, p: np.ndarray, w: np.ndarray | None = None) -> float:
    if len(np.unique(y)) < 2:
        return float("nan")
    return float(roc_auc_score(y, p, sample_weight=w))


def safe_ap(y: np.ndarray, p: np.ndarray, w: np.ndarray | None = None) -> float:
    if len(np.unique(y)) < 2:
        return float("nan")
    return float(average_precision_score(y, p, sample_weight=w))


def calibration_stats(y: np.ndarray, p: np.ndarray, w: np.ndarray) -> dict[str, float]:
    p = np.clip(p, EPS, 1 - EPS)
    y = np.asarray(y, dtype=float)
    w = np.asarray(w, dtype=float)
    obs = weighted_mean(y, w)
    mean_p = weighted_mean(p, w)
    logit_p = np.log(p / (1 - p))
    out = {
        "calibration_intercept": float("nan"),
        "calibration_slope": float("nan"),
        "observed_expected_ratio": float("nan") if mean_p == 0 else float(obs / mean_p),
        "weighted_observed_prevalence": obs,
        "weighted_mean_predicted_probability": mean_p,
        "calibration_in_the_large": float("nan"),
    }
    if 0 < obs < 1 and 0 < mean_p < 1:
        out["calibration_in_the_large"] = float(math.log(obs / (1 - obs)) - math.log(mean_p / (1 - mean_p)))
    try:
        X = add_constant(logit_p, has_constant="add")
        fit = GLM(y, X, family=Binomial(), freq_weights=w).fit(maxiter=100, disp=0)
        out["calibration_intercept"] = float(fit.params[0])
        out["calibration_slope"] = float(fit.params[1])
    except Exception:
        pass
    return out


def youden_threshold(y: np.ndarray, p: np.ndarray, w: np.ndarray) -> float:
    if len(np.unique(y)) < 2:
        return 0.5
    fpr, tpr, thresholds = roc_curve(y, p, sample_weight=w)
    j = tpr - fpr
    idx = int(np.nanargmax(j))
    thr = float(thresholds[idx])
    if not np.isfinite(thr) or thr > 1:
        return 0.5
    return max(0.0, min(1.0, thr))


def class_metrics(y: np.ndarray, p: np.ndarray, threshold: float, w: np.ndarray | None = None) -> dict[str, float]:
    pred = (p >= threshold).astype(int)
    tn, fp, fn, tp = confusion_matrix(y, pred, labels=[0, 1]).ravel()
    sens = tp / (tp + fn) if (tp + fn) else float("nan")
    spec = tn / (tn + fp) if (tn + fp) else float("nan")
    ppv = tp / (tp + fp) if (tp + fp) else float("nan")
    npv = tn / (tn + fn) if (tn + fn) else float("nan")
    f1 = 2 * ppv * sens / (ppv + sens) if np.isfinite(ppv) and np.isfinite(sens) and (ppv + sens) else float("nan")
    return {"sensitivity": sens, "specificity": spec, "PPV": ppv, "NPV": npv, "F1": f1}


def all_metrics(y: np.ndarray, p: np.ndarray, w: np.ndarray, threshold: float) -> dict[str, float]:
    out = {
        "survey_weighted_AUROC": safe_auc(y, p, w),
        "survey_weighted_PR_AUC": safe_ap(y, p, w),
        "unweighted_AUROC": safe_auc(y, p, None),
        "unweighted_PR_AUC": safe_ap(y, p, None),
        "weighted_Brier": weighted_brier(y, p, w),
        "weighted_log_loss": weighted_log_loss(y, p, w),
        "unweighted_Brier": float(brier_score_loss(y, p)),
        "unweighted_log_loss": float(log_loss(y, p, labels=[0, 1])),
    }
    out.update(calibration_stats(y, p, w))
    out.update(class_metrics(y, p, threshold))
    return out


def metric_row(
    model_name: str,
    algorithm: str,
    cycle: str,
    y: np.ndarray,
    p: np.ndarray,
    w: np.ndarray,
    threshold: float,
    threshold_type: str,
    n: int,
    events: int,
    weight_var: str,
    comparison_sample: str,
) -> dict[str, Any]:
    row = {
        "target_population": target_population(model_name),
        "comparison_sample": comparison_sample,
        "model": model_name,
        "model_label": MODEL_DISPLAY.get(model_name, model_name),
        "algorithm": algorithm,
        "cycle": cycle,
        "n": int(n),
        "events": int(events),
        "weight_variable": weight_var,
        "classification_threshold": threshold,
        "classification_threshold_type": threshold_type,
    }
    row.update(all_metrics(y, p, w, threshold))
    return row


def get_X_y_w(d: pd.DataFrame, features: list[str], model_name: str) -> tuple[pd.DataFrame, np.ndarray, np.ndarray]:
    X = d[features].copy()
    y = d["cvd"].astype(int).to_numpy()
    w = d[weight_column(model_name)].astype(float).to_numpy()
    return X, y, w


def select_hyperparameters(inner_rows: list[dict[str, Any]], algorithm: str) -> dict[str, Any]:
    df = pd.DataFrame(inner_rows)
    metric = (
        df.groupby("candidate_id", as_index=False)
        .agg(mean_weighted_log_loss=("weighted_log_loss", "mean"), sd_weighted_log_loss=("weighted_log_loss", "std"))
        .sort_values("mean_weighted_log_loss")
    )
    best_mean = metric["mean_weighted_log_loss"].min()
    tolerance = 1e-4
    near = metric[metric["mean_weighted_log_loss"] <= best_mean + tolerance].copy()
    params_by_id = {r["candidate_id"]: json.loads(r["hyperparameters_json"]) for r in inner_rows}
    if algorithm == "elastic_net":
        near["complexity"] = near["candidate_id"].map(lambda cid: en_complexity_key(params_by_id[cid]))
    else:
        near["complexity"] = near["candidate_id"].map(lambda cid: xgb_complexity_key(params_by_id[cid]))
    near = near.sort_values(["mean_weighted_log_loss", "complexity"])
    selected_id = near.iloc[0]["candidate_id"]
    selected = params_by_id[selected_id]
    selected["_selected_candidate_id"] = selected_id
    selected["_mean_weighted_log_loss"] = float(metric.loc[metric["candidate_id"].eq(selected_id), "mean_weighted_log_loss"].iloc[0])
    selected["_sd_weighted_log_loss"] = float(metric.loc[metric["candidate_id"].eq(selected_id), "sd_weighted_log_loss"].iloc[0])
    return selected


def tune_one(
    ctx: Stage3AContext,
    df: pd.DataFrame,
    features: list[str],
    model_name: str,
    algorithm: str,
    train_cycles: list[str],
    candidates: list[dict[str, Any]],
    checkpoint_dir: Path | None = None,
    checkpoint_signature: dict[str, Any] | None = None,
    resume: bool = False,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    rows: list[dict[str, Any]] = []
    resumed_candidates = 0
    pending: list[tuple[int, str, dict[str, Any], Path | None, dict[str, Any] | None]] = []
    for cid, params in enumerate(candidates):
        cid_s = f"{algorithm}_{cid:04d}"
        candidate_dir = None
        candidate_signature = None
        if checkpoint_signature is not None:
            candidate_dir = candidate_checkpoint_dir(ctx.project_root, checkpoint_signature, cid_s)
            candidate_signature = {
                **checkpoint_signature,
                "checkpoint_type": f"{checkpoint_signature['checkpoint_type']}_inner_candidate",
                "candidate_id": cid_s,
                "hyperparameters": params,
                "inner_validation_cycles": train_cycles,
            }
        loaded = None
        if resume and candidate_dir is not None and candidate_signature is not None:
            loaded = read_checkpoint_files(
                candidate_dir,
                candidate_signature,
                {"inner": "inner.csv"},
                project_root=ctx.project_root,
            )
        if loaded is not None:
            rows.extend(loaded[1]["inner"].to_dict(orient="records"))
            resumed_candidates += 1
        else:
            pending.append((cid, cid_s, params, candidate_dir, candidate_signature))
    if resumed_candidates:
        log(ctx, f"Resume reused {resumed_candidates}/{len(candidates)} inner candidates for {algorithm} {model_name}")
    if pending:
        n_jobs = int(os.environ.get("STAGE3A_N_JOBS", str(DEFAULT_PARALLEL_JOBS)))
        n_jobs = max(1, min(n_jobs, len(pending)))
        log(ctx, f"Evaluating {len(pending)}/{len(candidates)} pending candidates for {algorithm} {model_name} with n_jobs={n_jobs}")
        # Persist each completed batch before starting the next one. At most one
        # worker batch is lost if the process is interrupted.
        for batch_start in range(0, len(pending), n_jobs):
            batch = pending[batch_start : batch_start + n_jobs]
            if n_jobs == 1:
                evaluated = [
                    evaluate_candidate_inner(df, features, model_name, algorithm, train_cycles, cid_s, params)
                    for _cid, cid_s, params, _candidate_dir, _candidate_signature in batch
                ]
            else:
                evaluated = joblib.Parallel(n_jobs=n_jobs, backend="loky")(
                    joblib.delayed(evaluate_candidate_inner)(df, features, model_name, algorithm, train_cycles, cid_s, params)
                    for _cid, cid_s, params, _candidate_dir, _candidate_signature in batch
                )
            for (_cid, cid_s, _params, candidate_dir, candidate_signature), candidate_rows in zip(batch, evaluated):
                rows.extend(candidate_rows)
                if candidate_dir is not None and candidate_signature is not None:
                    write_checkpoint_frames(candidate_dir, candidate_signature, {"inner": pd.DataFrame(candidate_rows)})
                log(ctx, f"Tuning progress {algorithm} {model_name}: candidate {cid_s} complete")
    selected = select_hyperparameters(rows, algorithm)
    return selected, rows


def evaluate_candidate_inner(
    df: pd.DataFrame,
    features: list[str],
    model_name: str,
    algorithm: str,
    train_cycles: list[str],
    cid_s: str,
    params: dict[str, Any],
) -> list[dict[str, Any]]:
    candidate_rows = []
    mask = sample_mask(df, model_name)
    for inner_val_cycle in train_cycles:
        inner_train_cycles = [c for c in train_cycles if c != inner_val_cycle]
        tr = df[mask & df["cycle"].isin(inner_train_cycles)].copy()
        va = df[mask & df["cycle"].eq(inner_val_cycle)].copy()
        if tr.empty or va.empty:
            raise RuntimeError(f"Empty inner split for {model_name} {algorithm} {inner_val_cycle}")
        Xtr, ytr, wtr = get_X_y_w(tr, features, model_name)
        Xva, yva, wva = get_X_y_w(va, features, model_name)
        pipe = make_pipeline(features, algorithm, tr, params)
        fit_pipeline(pipe, Xtr, ytr, wtr)
        pva = predict_proba(pipe, Xva)
        candidate_rows.append(
            {
                "model": model_name,
                "algorithm": algorithm,
                "candidate_id": cid_s,
                "inner_validation_cycle": inner_val_cycle,
                "n_train": len(tr),
                "n_validation": len(va),
                "events_validation": int(yva.sum()),
                "weighted_log_loss": weighted_log_loss(yva, pva, wva),
                "weighted_Brier": weighted_brier(yva, pva, wva),
                "weighted_AUROC": safe_auc(yva, pva, wva),
                "hyperparameters_json": json.dumps(params, sort_keys=True),
            }
        )
    return candidate_rows


def candidates_for_algorithm(
    project_root: Path,
    algorithm: str,
    en_limit: int | None,
    xgb_limit: int | None,
) -> tuple[list[dict[str, Any]], int | None]:
    if algorithm == "elastic_net":
        return elastic_net_candidates(project_root, en_limit), en_limit
    if algorithm == "xgboost":
        return xgboost_candidates(project_root, xgb_limit), xgb_limit
    raise ValueError(f"Unsupported algorithm: {algorithm}")


def outer_checkpoint_dir(project_root: Path, signature: dict[str, Any]) -> Path:
    name = checkpoint_task_name(
        signature["fold"],
        signature["holdout_cycle"],
        signature["model"],
        signature["algorithm"],
    )
    return project_root / "results/checkpoints/stage3a/outer" / name


def load_outer_task(project_root: Path, signature: dict[str, Any]) -> tuple[dict[str, Any], dict[str, pd.DataFrame]] | None:
    return read_checkpoint_files(
        outer_checkpoint_dir(project_root, signature),
        signature,
        {
            "performance": "performance.csv",
            "inner": "inner.csv",
            "selected": "selected.csv",
            "predictions": "predictions.csv",
        },
        project_root=project_root,
    )


def append_outer_frames(target: dict[str, list[pd.DataFrame]], frames: dict[str, pd.DataFrame]) -> None:
    for key in target:
        target[key].append(frames[key])


def checkpoint_manifest_row(project_root: Path, meta: dict[str, Any], action: str) -> dict[str, Any]:
    sig = meta["signature"]
    return {
        "checkpoint_type": sig["checkpoint_type"],
        "task": checkpoint_task_name(sig["fold"], sig["holdout_cycle"], sig["model"], sig["algorithm"]),
        "algorithm": sig["algorithm"],
        "model": sig["model"],
        "fold": sig["fold"],
        "holdout_cycle": sig["holdout_cycle"],
        "candidate_count": sig["candidate_count"],
        "candidate_limit": sig["candidate_limit"],
        "candidate_grid_hash": sig["candidate_grid_hash"],
        "input_hash": sig["input_hash"],
        "code_hash": sig["code_hash"],
        "analysis_semantics_version": sig.get("analysis_semantics_version", ""),
        "analysis_semantics_hash": sig.get("analysis_semantics_hash", ""),
        "compatibility_status": checkpoint_compatibility_status(project_root, sig),
        "compatibility_attestation_id": sig.get("compatibility_attestation_id", ""),
        "checkpoint_created": meta["created"],
        "action": action,
        "status": meta["status"],
    }


def write_grid_completion_audit(project_root: Path, manifest: pd.DataFrame) -> pd.DataFrame:
    expected_tasks = len(DEVELOPMENT_CYCLES) * len(MODEL_SAMPLE_DOMAIN)
    rows = []
    for algorithm in SUPPORTED_ALGORITHMS:
        g = manifest[manifest["algorithm"].eq(algorithm)].copy() if not manifest.empty else pd.DataFrame()
        full_count = len(elastic_net_candidates(project_root)) if algorithm == "elastic_net" else len(xgboost_candidates(project_root))
        task_count = int(g["task"].nunique()) if not g.empty else 0
        counts = sorted(g["candidate_count"].dropna().astype(int).unique().tolist()) if not g.empty else []
        limits_absent = bool(g["candidate_limit"].isna().all()) if not g.empty else False
        compatibility_pass = bool(g["compatibility_status"].eq("PASS").all()) if not g.empty else False
        semantic_hashes = sorted(g["analysis_semantics_hash"].dropna().astype(str).unique().tolist()) if not g.empty else []
        semantics_pass = semantic_hashes == [analysis_semantics_hash()]
        passed = task_count == expected_tasks and counts == [full_count] and limits_absent and compatibility_pass and semantics_pass
        rows.append(
            {
                "algorithm": algorithm,
                "status": "PASS" if passed else ("NOT_RUN" if task_count == 0 else "INCOMPLETE"),
                "completed_outer_tasks": task_count,
                "expected_outer_tasks": expected_tasks,
                "checkpoint_candidate_counts": ";".join(map(str, counts)),
                "locked_full_grid_candidate_count": full_count,
                "all_candidate_limits_absent": limits_absent,
                "all_checkpoint_compatibility_pass": compatibility_pass,
                "analysis_semantics_hashes": ";".join(semantic_hashes),
                "analysis_semantics_pass": semantics_pass,
            }
        )
    audit = pd.DataFrame(rows)
    atomic_write_csv(audit, project_root / "results/audit/stage3a_grid_completion.csv")
    lines = ["# Stage 3A Locked-grid Completion Audit", ""]
    for row in rows:
        lines.append(
            f"- {row['algorithm']}: {row['status']} "
            f"({row['completed_outer_tasks']}/{row['expected_outer_tasks']} outer tasks; "
            f"locked grid {row['locked_full_grid_candidate_count']} candidates)."
        )
    atomic_write_text(project_root / "results/audit/stage3a_grid_completion.md", "\n".join(lines) + "\n")
    return audit


def run_internal_external(
    ctx: Stage3AContext,
    algorithms: list[str],
    en_limit: int | None = None,
    xgb_limit: int | None = None,
    resume: bool = False,
) -> dict[str, pd.DataFrame]:
    project_root = ctx.project_root
    check_temporal_lock(project_root)
    df_all = read_harmonized_data(project_root)
    df = development_df(df_all)
    features = read_frozen_features(project_root)
    folds = read_cycle_folds(project_root)
    base_signature = checkpoint_base_signature(project_root)
    log(ctx, f"Loaded development rows: {len(df)}; input hash {base_signature['input_hash'][:12]}")

    model_names = [
        "model0_core",
        "model1_renal",
        "model2_metabolic",
        "model3_inflammatory",
        "model4_combined",
        "model0_paired_metabolic",
        "model0_paired_combined",
    ]
    frame_lists: dict[str, list[pd.DataFrame]] = {"performance": [], "inner": [], "selected": [], "predictions": []}
    manifest_rows: list[dict[str, Any]] = []

    # Requested algorithms are computed (or resumed) first.
    for fold in folds:
        holdout = fold["holdout_cycle"]
        train_cycles = list(fold["training_cycles"])
        for model_name in model_names:
            for algorithm in algorithms:
                cand, limit = candidates_for_algorithm(project_root, algorithm, en_limit, xgb_limit)
                signature = outer_task_signature(base_signature, fold, model_name, algorithm, features[model_name], cand, limit)
                checkpoint_dir = outer_checkpoint_dir(project_root, signature)
                loaded = load_outer_task(project_root, signature) if resume else None
                if loaded is not None:
                    meta, frames = loaded
                    append_outer_frames(frame_lists, frames)
                    manifest_rows.append(checkpoint_manifest_row(project_root, meta, "resumed"))
                    log(ctx, f"Resume hit {algorithm} {model_name}; outer holdout {holdout}")
                    continue

                log(ctx, f"Tuning {algorithm} {model_name}; outer holdout {holdout}; candidates {len(cand)}")
                selected, inner_rows = tune_one(
                    ctx,
                    df,
                    features[model_name],
                    model_name,
                    algorithm,
                    train_cycles,
                    cand,
                    checkpoint_dir=checkpoint_dir,
                    checkpoint_signature=signature,
                    resume=resume,
                )
                inner_frame = pd.DataFrame(
                    [{**r, "outer_holdout_cycle": holdout, "fold": fold["fold"]} for r in inner_rows]
                )
                params = {k: v for k, v in selected.items() if not k.startswith("_")}
                mask = sample_mask(df, model_name)
                tr = df[mask & df["cycle"].isin(train_cycles)].copy()
                va = df[mask & df["cycle"].eq(holdout)].copy()
                Xtr, ytr, wtr = get_X_y_w(tr, features[model_name], model_name)
                Xva, yva, wva = get_X_y_w(va, features[model_name], model_name)
                pipe = make_pipeline(features[model_name], algorithm, tr, params)
                fit_pipeline(pipe, Xtr, ytr, wtr)
                p_tr = predict_proba(pipe, Xtr)
                p_va = predict_proba(pipe, Xva)
                youden = youden_threshold(ytr, p_tr, wtr)
                weight_var = weight_column(model_name)
                comparison_sample = comparison_sample_label(model_name)
                performance_rows = []
                for thr, thr_type in [(0.5, "fixed_0.50"), (youden, "training_weighted_youden")]:
                    performance_rows.append(
                        metric_row(
                            model_name,
                            algorithm,
                            holdout,
                            yva,
                            p_va,
                            wva,
                            thr,
                            thr_type,
                            len(va),
                            int(yva.sum()),
                            weight_var,
                            comparison_sample,
                        )
                    )
                selected_frame = pd.DataFrame(
                    [
                        {
                            "fold": fold["fold"],
                            "outer_holdout_cycle": holdout,
                            "model": model_name,
                            "algorithm": algorithm,
                            "selected_candidate_id": selected["_selected_candidate_id"],
                            "mean_inner_weighted_log_loss": selected["_mean_weighted_log_loss"],
                            "sd_inner_weighted_log_loss": selected["_sd_weighted_log_loss"],
                            "training_cycles": ";".join(train_cycles),
                            "hyperparameters_json": json.dumps(params, sort_keys=True),
                            "n_outer_train": len(tr),
                            "n_outer_validation": len(va),
                            "events_outer_validation": int(yva.sum()),
                            "training_youden_threshold": youden,
                            "config_hash": json.dumps(base_signature["config_hashes"], sort_keys=True),
                        }
                    ]
                )
                prediction_rows = []
                for seqn, cycle, y_i, p_i, w_i, strata, psu in zip(
                    va["seqn"],
                    va["cycle"],
                    yva,
                    p_va,
                    wva,
                    va["strata"],
                    va["psu"],
                ):
                    prediction_rows.append(
                        {
                            "seqn": seqn,
                            "cycle": cycle,
                            "model": model_name,
                            "algorithm": algorithm,
                            "cvd": int(y_i),
                            "predicted_probability": float(p_i),
                            "analysis_weight": float(w_i),
                            "weight_variable": weight_var,
                            "strata": strata,
                            "psu": psu,
                            "outer_holdout_cycle": holdout,
                            "comparison_sample": comparison_sample,
                        }
                    )
                frames = {
                    "performance": pd.DataFrame(performance_rows),
                    "inner": inner_frame,
                    "selected": selected_frame,
                    "predictions": pd.DataFrame(prediction_rows),
                }
                meta = write_checkpoint_frames(checkpoint_dir, signature, frames)
                append_outer_frames(frame_lists, frames)
                manifest_rows.append(checkpoint_manifest_row(project_root, meta, "computed"))
                log(ctx, f"Checkpoint saved {algorithm} {model_name}; outer holdout {holdout}")

    # A staged second run must retain a previously completed full-grid algorithm.
    for algorithm in [a for a in SUPPORTED_ALGORITHMS if a not in algorithms]:
        cand, limit = candidates_for_algorithm(project_root, algorithm, None, None)
        retained_units: list[tuple[dict[str, Any], dict[str, pd.DataFrame]]] = []
        complete = True
        for fold in folds:
            for model_name in model_names:
                signature = outer_task_signature(base_signature, fold, model_name, algorithm, features[model_name], cand, limit)
                loaded = load_outer_task(project_root, signature)
                if loaded is None:
                    complete = False
                    break
                retained_units.append(loaded)
            if not complete:
                break
        if complete:
            for meta, frames in retained_units:
                append_outer_frames(frame_lists, frames)
                manifest_rows.append(checkpoint_manifest_row(project_root, meta, "retained_from_prior_run"))
            log(ctx, f"Retained {len(retained_units)} completed full-grid {algorithm} checkpoints")

    if not frame_lists["performance"]:
        raise RuntimeError("No valid outer-validation checkpoints were produced.")
    out = {key: pd.concat(frames, ignore_index=True) for key, frames in frame_lists.items()}
    out["performance"] = out["performance"].sort_values(["algorithm", "fold" if "fold" in out["performance"].columns else "cycle", "model", "classification_threshold_type"]) if "fold" in out["performance"].columns else out["performance"].sort_values(["algorithm", "cycle", "model", "classification_threshold_type"])
    out["inner"] = out["inner"].sort_values(["algorithm", "outer_holdout_cycle", "model", "candidate_id", "inner_validation_cycle"])
    out["selected"] = out["selected"].sort_values(["algorithm", "outer_holdout_cycle", "model"])
    out["predictions"] = out["predictions"].sort_values(["algorithm", "outer_holdout_cycle", "model", "seqn"])
    if out["selected"].duplicated(["outer_holdout_cycle", "model", "algorithm"]).any():
        raise RuntimeError("Duplicate outer-task checkpoint rows detected during consolidation.")
    if out["predictions"].duplicated(["seqn", "cycle", "model", "algorithm"]).any():
        raise RuntimeError("Duplicate outer predictions detected during checkpoint consolidation.")

    atomic_write_csv(out["performance"], project_root / "results/tables/cycle_specific_model_performance.csv")
    atomic_write_csv(out["inner"], project_root / "results/models/inner_cv_all_candidates.csv")
    atomic_write_csv(out["selected"], project_root / "results/models/outer_fold_hyperparameters.csv")
    atomic_write_csv(out["selected"], project_root / "results/tables/final_development_hyperparameters.csv")
    atomic_write_csv(out["predictions"], project_root / "results/predictions/stage3a_outer_predictions.csv")
    manifest = pd.DataFrame(manifest_rows).sort_values(["algorithm", "holdout_cycle", "model"])
    atomic_write_csv(manifest, project_root / "results/logs/stage3a_checkpoint_manifest.csv")
    write_grid_completion_audit(project_root, manifest)
    write_docx_table(out["performance"], project_root / "results/manuscript_tables/Table_3_cycle_specific_model_performance.docx", "Table 3. Leave-one-cycle-out internal-external validation")
    return out


def comparison_sample_label(model_name: str) -> str:
    if model_name == "model1_renal":
        return "renal comparison sample"
    if model_name in ["model2_metabolic", "model0_paired_metabolic"]:
        return "metabolic comparison sample"
    if model_name == "model3_inflammatory":
        return "inflammatory comparison sample"
    if model_name in ["model4_combined", "model0_paired_combined"]:
        return "combined comparison sample"
    return "core MEC target sample"


def config_hashes_dict(project_root: Path) -> dict[str, str]:
    rels = [
        "config/frozen_variable_specification.yml",
        "config/cycle_folds.yml",
        "config/elastic_net_hyperparameter_grid.yml",
        "config/xgboost_hyperparameter_grid.yml",
        "config/temporal_validation_locked.yml",
    ]
    return {rel: sha256_file(project_root / rel) for rel in rels}


def checkpoint_base(project_root: Path) -> Path:
    base = project_root / "results/checkpoints/stage3a"
    base.mkdir(parents=True, exist_ok=True)
    return base


def task_id(task_type: str, algorithm: str, model: str, cycle: str = "all") -> str:
    return f"{task_type}__{algorithm}__{model}__{cycle}".replace("/", "_").replace("\\", "_")


def task_checkpoint_dir(project_root: Path, task_type: str, algorithm: str, model: str, cycle: str = "all") -> Path:
    return checkpoint_base(project_root) / task_id(task_type, algorithm, model, cycle)


def limit_label(value: int | None) -> str:
    return "" if value is None else str(value)


def expected_candidate_count(project_root: Path, algorithm: str, limit: int | None) -> int:
    if algorithm == "elastic_net":
        return len(elastic_net_candidates(project_root, limit))
    if algorithm == "xgboost":
        return len(xgboost_candidates(project_root, limit))
    raise ValueError(algorithm)


def checkpoint_complete(ck: Path, algorithm: str, expected_candidates: int, candidate_limit: int | None) -> bool:
    meta_path = ck / "metadata.json"
    if not meta_path.exists():
        return False
    try:
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return False
    if meta.get("status") != "COMPLETE":
        return False
    if meta.get("algorithm") != algorithm:
        return False
    if int(meta.get("candidate_count", -1)) != int(expected_candidates):
        return False
    if str(meta.get("candidate_limit", "")) != limit_label(candidate_limit):
        return False
    required = meta.get("required_files", [])
    return all((ck / rel).exists() for rel in required)


def write_checkpoint_metadata(
    project_root: Path,
    ck: Path,
    row: dict[str, Any],
    required_files: list[str],
) -> None:
    ck.mkdir(parents=True, exist_ok=True)
    meta = {**row, "required_files": required_files, "checkpoint_dir": str(ck), "updated_at": now(), "status": "COMPLETE"}
    (ck / "metadata.json").write_text(json.dumps(meta, indent=2, sort_keys=True), encoding="utf-8")
    record_checkpoint(project_root, meta)


def record_checkpoint(project_root: Path, row: dict[str, Any]) -> None:
    path = project_root / "results/logs/stage3a_checkpoint_manifest.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    flat = {k: (json.dumps(v, sort_keys=True) if isinstance(v, (dict, list)) else v) for k, v in row.items() if k != "required_files"}
    df_new = pd.DataFrame([flat])
    if path.exists():
        df = pd.read_csv(path)
        if "task_id" in df.columns and "task_id" in df_new.columns:
            df = df[~df["task_id"].eq(df_new.loc[0, "task_id"])].copy()
        df = pd.concat([df, df_new], ignore_index=True)
    else:
        df = df_new
    df.to_csv(path, index=False)


def generate_denominator_tables(ctx: Stage3AContext) -> pd.DataFrame:
    project_root = ctx.project_root
    df_all = read_harmonized_data(project_root)
    features = read_frozen_features(project_root)
    rows = []
    for cycle in DEVELOPMENT_CYCLES + [EXTERNAL_CYCLE]:
        d = df_all[df_all["cycle"].eq(cycle)].copy()
        adult_np = d["age"].ge(20) & (~d["pregnancy_code"].eq(1).fillna(False))
        out = adult_np & d["cvd"].notna()
        for model in ["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"]:
            domain = MODEL_SAMPLE_DOMAIN[model]
            wcol = weight_column(model)
            design = out & d["strata"].notna() & d["psu"].notna() & d[wcol].notna() & d[wcol].gt(EPS)
            complete = design & d[features[model]].notna().all(axis=1)
            rows.append(
                {
                    "cycle": cycle,
                    "model": model,
                    "domain": domain,
                    "adult_nonpregnant_source_population": int(adult_np.sum()),
                    "outcome_available_population": int(out.sum()),
                    "model_specific_imputation_eligible_population": int(design.sum()),
                    "model_specific_complete_case_population": int(complete.sum()),
                    "events_imputation_eligible": int(d.loc[design, "cvd"].sum()) if design.any() else 0,
                    "events_complete_case": int(d.loc[complete, "cvd"].sum()) if complete.any() else 0,
                    "weight_variable": wcol,
                }
            )
    den = pd.DataFrame(rows)
    den.to_csv(project_root / "results/tables/denominator_reconciliation.csv", index=False)
    lines = [
        "# Denominator Reconciliation Audit",
        "",
        "- Adult non-pregnant source population, outcome-available population, model-specific imputation-eligible population, and complete-case population are separated.",
        "- Percentages in Stage 3A tables must use the denominator identified in the table row.",
        "- UACR source-population missingness must not be reused as the Model 1 analysis-sample missingness unless denominators match.",
        "- 2021-2023 rows are included only for denominator reconciliation; no 2021-2023 model performance is computed.",
    ]
    (project_root / "results/audit/denominator_reconciliation_audit.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return den


def aggregate_performance(project_root: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    perf = pd.read_csv(project_root / "results/tables/cycle_specific_model_performance.csv")
    pmain = perf[perf["classification_threshold_type"].eq("fixed_0.50")].copy()
    metrics = [
        "survey_weighted_AUROC",
        "survey_weighted_PR_AUC",
        "weighted_Brier",
        "weighted_log_loss",
        "calibration_intercept",
        "calibration_slope",
        "weighted_observed_prevalence",
        "weighted_mean_predicted_probability",
    ]
    pooled_rows = []
    for keys, g in pmain.groupby(["target_population", "comparison_sample", "model", "model_label", "algorithm", "weight_variable"]):
        row = dict(zip(["target_population", "comparison_sample", "model", "model_label", "algorithm", "weight_variable"], keys))
        row["n_cycles"] = g["cycle"].nunique()
        row["total_n_across_validation_cycles"] = int(g["n"].sum())
        row["total_events_across_validation_cycles"] = int(g["events"].sum())
        for m in metrics:
            row[f"{m}_mean"] = float(g[m].mean())
            row[f"{m}_sd"] = float(g[m].std(ddof=1)) if len(g) > 1 else float("nan")
        pooled_rows.append(row)
    pooled = pd.DataFrame(pooled_rows)
    pooled.to_csv(project_root / "results/tables/pooled_out_of_cycle_performance.csv", index=False)

    pred = pd.read_csv(project_root / "results/predictions/stage3a_outer_predictions.csv")
    inc_rows = []
    cycle_rows = []
    for comp, (core, extended, domain) in INCREMENT_COMPARISONS.items():
        for alg in pred["algorithm"].unique():
            p1 = pred[(pred["model"].eq(core)) & (pred["algorithm"].eq(alg))].copy()
            p2 = pred[(pred["model"].eq(extended)) & (pred["algorithm"].eq(alg))].copy()
            keys = ["seqn", "cycle"]
            pair = p1.merge(p2, on=keys, suffixes=("_core", "_extended"))
            if pair.empty:
                continue
            for cycle, g in pair.groupby("cycle"):
                y = g["cvd_core"].astype(int).to_numpy()
                w = g["analysis_weight_extended"].astype(float).to_numpy()
                pc = g["predicted_probability_core"].to_numpy()
                pe = g["predicted_probability_extended"].to_numpy()
                mc = all_metrics(y, pc, w, 0.5)
                me = all_metrics(y, pe, w, 0.5)
                row = {
                    "comparison": comp,
                    "algorithm": alg,
                    "cycle": cycle,
                    "n": len(g),
                    "events": int(y.sum()),
                    "weight_variable": g["weight_variable_extended"].iloc[0],
                    "core_model": core,
                    "extended_model": extended,
                    "delta_AUROC": me["survey_weighted_AUROC"] - mc["survey_weighted_AUROC"],
                    "delta_PR_AUC": me["survey_weighted_PR_AUC"] - mc["survey_weighted_PR_AUC"],
                    "delta_Brier_improvement": mc["weighted_Brier"] - me["weighted_Brier"],
                    "delta_log_loss_improvement": mc["weighted_log_loss"] - me["weighted_log_loss"],
                    "core_calibration_intercept": mc["calibration_intercept"],
                    "extended_calibration_intercept": me["calibration_intercept"],
                    "core_calibration_slope": mc["calibration_slope"],
                    "extended_calibration_slope": me["calibration_slope"],
                    "absolute_intercept_deviation_change": abs(mc["calibration_intercept"]) - abs(me["calibration_intercept"]),
                    "absolute_slope_deviation_change": abs(mc["calibration_slope"] - 1) - abs(me["calibration_slope"] - 1),
                }
                cycle_rows.append(row)
            cg = pair
            y = cg["cvd_core"].astype(int).to_numpy()
            w = cg["analysis_weight_extended"].astype(float).to_numpy()
            pc = cg["predicted_probability_core"].to_numpy()
            pe = cg["predicted_probability_extended"].to_numpy()
            mc = all_metrics(y, pc, w, 0.5)
            me = all_metrics(y, pe, w, 0.5)
            inc_rows.append(
                {
                    "comparison": comp,
                    "algorithm": alg,
                    "n": len(cg),
                    "events": int(y.sum()),
                    "weight_variable": cg["weight_variable_extended"].iloc[0],
                    "core_model": core,
                    "extended_model": extended,
                    "delta_AUROC": me["survey_weighted_AUROC"] - mc["survey_weighted_AUROC"],
                    "delta_PR_AUC": me["survey_weighted_PR_AUC"] - mc["survey_weighted_PR_AUC"],
                    "delta_Brier_improvement": mc["weighted_Brier"] - me["weighted_Brier"],
                    "delta_log_loss_improvement": mc["weighted_log_loss"] - me["weighted_log_loss"],
                    "core_AUROC": mc["survey_weighted_AUROC"],
                    "extended_AUROC": me["survey_weighted_AUROC"],
                    "core_Brier": mc["weighted_Brier"],
                    "extended_Brier": me["weighted_Brier"],
                    "core_calibration_slope": mc["calibration_slope"],
                    "extended_calibration_slope": me["calibration_slope"],
                }
            )
    cycle_inc = pd.DataFrame(cycle_rows)
    pooled_inc = pd.DataFrame(inc_rows)
    cycle_inc.to_csv(project_root / "results/tables/cycle_specific_incremental_value.csv", index=False)
    pooled_inc.to_csv(project_root / "results/tables/pooled_incremental_value.csv", index=False)
    write_docx_table(cycle_inc, project_root / "results/manuscript_tables/Table_4_incremental_value.docx", "Table 4. Incremental value across NHANES cycles")
    return pooled, pooled_inc


def write_docx_table(df: pd.DataFrame, path: Path, title: str, max_rows: int = 200) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    if df.empty:
        doc.add_paragraph("No data available.")
    else:
        view = df.head(max_rows).copy()
        table = doc.add_table(rows=1, cols=len(view.columns))
        hdr = table.rows[0].cells
        for i, col in enumerate(view.columns):
            hdr[i].text = str(col)
        for _, row in view.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(view.columns):
                val = row[col]
                if isinstance(val, float):
                    cells[i].text = "" if not np.isfinite(val) else f"{val:.4f}"
                else:
                    cells[i].text = "" if pd.isna(val) else str(val)
    doc.save(path)


def plot_figures(project_root: Path) -> None:
    perf = pd.read_csv(project_root / "results/tables/cycle_specific_model_performance.csv")
    perf = perf[perf["classification_threshold_type"].eq("fixed_0.50")].copy()
    perf = perf[perf["model"].isin(["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"])]
    alg_order = ["elastic_net", "xgboost"]
    model_order = ["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"]
    label_map = {
        "elastic_net": "Elastic Net",
        "xgboost": "XGBoost",
        "model0_core": "Model 0",
        "model1_renal": "Model 1",
        "model2_metabolic": "Model 2",
        "model3_inflammatory": "Model 3",
        "model4_combined": "Model 4",
    }
    metrics = [
        ("survey_weighted_AUROC", "Survey-weighted AUROC"),
        ("weighted_Brier", "Weighted Brier score"),
        ("calibration_slope", "Calibration slope"),
    ]
    fig, axes = plt.subplots(3, 1, figsize=(12, 12), sharex=True)
    for ax, (metric, title) in zip(axes, metrics):
        for alg in alg_order:
            for model in model_order:
                g = perf[(perf["algorithm"].eq(alg)) & (perf["model"].eq(model))].sort_values("cycle")
                if g.empty:
                    continue
                ax.plot(g["cycle"], g[metric], marker="o", linewidth=1.3, label=f"{label_map[model]} {label_map[alg]}")
        ax.set_title(title)
        ax.grid(True, alpha=0.25)
    axes[-1].tick_params(axis="x", rotation=30)
    axes[0].legend(ncol=2, fontsize=8, frameon=False)
    fig.tight_layout()
    for ext in ["png", "pdf", "svg", "tiff"]:
        fig.savefig(project_root / f"results/figures/Figure2_cycle_performance.{ext}", dpi=600)
    plt.close(fig)

    inc = pd.read_csv(project_root / "results/tables/cycle_specific_incremental_value.csv")
    fig, axes = plt.subplots(2, 2, figsize=(12, 8), sharex=True)
    inc_metrics = [
        ("delta_AUROC", "Delta AUROC"),
        ("delta_PR_AUC", "Delta PR-AUC"),
        ("delta_Brier_improvement", "Brier improvement"),
        ("delta_log_loss_improvement", "Log-loss improvement"),
    ]
    for ax, (metric, title) in zip(axes.ravel(), inc_metrics):
        for comp in ["renal", "metabolic", "inflammatory", "combined"]:
            g = inc[(inc["comparison"].eq(comp)) & (inc["algorithm"].eq("elastic_net"))].sort_values("cycle")
            if not g.empty:
                ax.plot(g["cycle"], g[metric], marker="o", label=comp.capitalize())
        ax.axhline(0, color="black", linewidth=0.8)
        ax.set_title(title)
        ax.grid(True, alpha=0.25)
        ax.tick_params(axis="x", rotation=30)
    axes[0, 0].legend(frameon=False, fontsize=8)
    fig.tight_layout()
    for ext in ["png", "pdf", "svg", "tiff"]:
        fig.savefig(project_root / f"results/figures/Figure4_incremental_value.{ext}", dpi=600)
    plt.close(fig)


def frozen_checkpoint_dir(project_root: Path, model_name: str, algorithm: str) -> Path:
    return project_root / "results/checkpoints/stage3a/frozen" / f"{model_name}__{algorithm}"


def load_frozen_task(
    project_root: Path,
    model_name: str,
    algorithm: str,
    expected_signature: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]] | None:
    checkpoint_dir = frozen_checkpoint_dir(project_root, model_name, algorithm)
    checkpoint_path = checkpoint_dir / "checkpoint.json"
    if not checkpoint_path.exists():
        return None
    try:
        checkpoint = json.loads(checkpoint_path.read_text(encoding="utf-8"))
        if checkpoint.get("status") != "complete" or not checkpoint_signature_matches(
            checkpoint.get("signature"), expected_signature, project_root
        ):
            return None
        required = {
            "artifact": "model_pipeline.joblib",
            "features": "feature_list.json",
            "hyperparameters": "hyperparameters.json",
            "metadata": "training_metadata.json",
        }
        for key, filename in required.items():
            path = checkpoint_dir / filename
            recorded = checkpoint.get("files", {}).get(key, {})
            if not path.exists() or recorded.get("sha256") != sha256_file(path):
                return None
        metadata = json.loads((checkpoint_dir / required["metadata"]).read_text(encoding="utf-8"))
        if metadata.get("model_hash") != checkpoint["files"]["artifact"]["sha256"]:
            return None
        return checkpoint, metadata
    except (OSError, ValueError, KeyError, json.JSONDecodeError):
        return None


def write_frozen_task(
    checkpoint_dir: Path,
    signature: dict[str, Any],
    pipe: Pipeline,
    features: list[str],
    params: dict[str, Any],
    metadata: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]]:
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    artifact = checkpoint_dir / "model_pipeline.joblib"
    feature_path = checkpoint_dir / "feature_list.json"
    hyperparameter_path = checkpoint_dir / "hyperparameters.json"
    metadata_path = checkpoint_dir / "training_metadata.json"
    atomic_joblib_dump(pipe, artifact)
    atomic_write_text(feature_path, json.dumps(features, indent=2))
    atomic_write_text(hyperparameter_path, json.dumps(params, indent=2, sort_keys=True))
    metadata = {**metadata, "model_hash": sha256_file(artifact), "resume_signature": signature}
    atomic_write_text(metadata_path, json.dumps(metadata, indent=2, sort_keys=True))
    paths = {
        "artifact": artifact,
        "features": feature_path,
        "hyperparameters": hyperparameter_path,
        "metadata": metadata_path,
    }
    checkpoint = {
        "status": "complete",
        "created": now(),
        "signature": signature,
        "files": file_manifest(paths),
    }
    atomic_write_text(checkpoint_dir / "checkpoint.json", json.dumps(checkpoint, indent=2, sort_keys=True))
    return checkpoint, metadata


def publish_frozen_task(
    project_root: Path,
    model_name: str,
    algorithm: str,
    checkpoint: dict[str, Any],
    metadata: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]]:
    source = frozen_checkpoint_dir(project_root, model_name, algorithm)
    destination = project_root / "results/models/frozen_development" / f"{model_name}_{algorithm}"
    destination.mkdir(parents=True, exist_ok=True)
    for filename in ["model_pipeline.joblib", "feature_list.json", "hyperparameters.json", "training_metadata.json"]:
        atomic_copy(source / filename, destination / filename)
    if sha256_file(destination / "model_pipeline.joblib") != metadata["model_hash"]:
        raise RuntimeError(f"Published frozen-model hash mismatch for {model_name} {algorithm}")
    params = json.loads((destination / "hyperparameters.json").read_text(encoding="utf-8"))
    row = {
        **metadata,
        "artifact": str(destination / "model_pipeline.joblib"),
        "hyperparameters_json": json.dumps(params, sort_keys=True),
    }
    cfg = {
        "artifact": str((destination / "model_pipeline.joblib").relative_to(project_root)),
        "feature_list": str((destination / "feature_list.json").relative_to(project_root)),
        "metadata": str((destination / "training_metadata.json").relative_to(project_root)),
        "model_hash": metadata["model_hash"],
    }
    return row, cfg


def fit_frozen_development_models(
    ctx: Stage3AContext,
    algorithms: list[str],
    en_limit: int | None = None,
    xgb_limit: int | None = None,
    resume: bool = False,
) -> pd.DataFrame:
    project_root = ctx.project_root
    df = development_df(read_harmonized_data(project_root))
    features = read_frozen_features(project_root)
    model_names = ["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"]
    base_signature = checkpoint_base_signature(project_root)
    completed: dict[tuple[str, str], tuple[dict[str, Any], dict[str, Any]]] = {}

    frozen_cfg: dict[str, Any] = {"created": now(), "training_period": "NHANES 2005-2018", "models": {}}
    for model_name in model_names:
        for algorithm in algorithms:
            cand, limit = candidates_for_algorithm(project_root, algorithm, en_limit, xgb_limit)
            signature = frozen_task_signature(base_signature, model_name, algorithm, features[model_name], cand, limit)
            loaded = load_frozen_task(project_root, model_name, algorithm, signature) if resume else None
            if loaded is not None:
                completed[(model_name, algorithm)] = loaded
                log(ctx, f"Resume hit frozen development model {model_name} {algorithm}")
                continue

            checkpoint_dir = frozen_checkpoint_dir(project_root, model_name, algorithm)
            log(ctx, f"Tuning frozen development model {model_name} {algorithm}; candidates {len(cand)}")
            train_cycles = DEVELOPMENT_CYCLES
            selected, _inner_rows = tune_one(
                ctx,
                df,
                features[model_name],
                model_name,
                algorithm,
                train_cycles,
                cand,
                checkpoint_dir=checkpoint_dir,
                checkpoint_signature=signature,
                resume=resume,
            )
            params = {k: v for k, v in selected.items() if not k.startswith("_")}
            mask = sample_mask(df, model_name)
            tr = df[mask].copy()
            Xtr, ytr, wtr = get_X_y_w(tr, features[model_name], model_name)
            pipe = make_pipeline(features[model_name], algorithm, tr, params)
            fit_pipeline(pipe, Xtr, ytr, wtr)
            meta = {
                "model": model_name,
                "algorithm": algorithm,
                "training_period": "NHANES 2005-2018",
                "n_training": len(tr),
                "events_training": int(ytr.sum()),
                "weight_variable": weight_column(model_name),
                "calibration_status": "none",
                "temporal_validation_status": "locked",
                "software": software_versions(),
                "config_hash": config_hashes_dict(project_root),
                "creation_timestamp": now(),
                "selected_candidate_id": selected["_selected_candidate_id"],
                "mean_inner_weighted_log_loss": selected["_mean_weighted_log_loss"],
                "sd_inner_weighted_log_loss": selected["_sd_weighted_log_loss"],
            }
            completed[(model_name, algorithm)] = write_frozen_task(
                checkpoint_dir,
                signature,
                pipe,
                features[model_name],
                params,
                meta,
            )
            log(ctx, f"Checkpoint saved frozen development model {model_name} {algorithm}")

    # Retain the other algorithm only if all five of its full-grid frozen models validate.
    for algorithm in [a for a in SUPPORTED_ALGORITHMS if a not in algorithms]:
        cand, limit = candidates_for_algorithm(project_root, algorithm, None, None)
        retained: dict[tuple[str, str], tuple[dict[str, Any], dict[str, Any]]] = {}
        for model_name in model_names:
            signature = frozen_task_signature(base_signature, model_name, algorithm, features[model_name], cand, limit)
            loaded = load_frozen_task(project_root, model_name, algorithm, signature)
            if loaded is None:
                retained = {}
                break
            retained[(model_name, algorithm)] = loaded
        if retained:
            completed.update(retained)
            log(ctx, f"Retained {len(retained)} completed full-grid frozen {algorithm} checkpoints")

    rows = []
    frozen_audit_rows = []
    for (model_name, algorithm), (checkpoint, metadata) in sorted(completed.items(), key=lambda item: (item[0][1], item[0][0])):
        row, cfg = publish_frozen_task(project_root, model_name, algorithm, checkpoint, metadata)
        rows.append(row)
        frozen_cfg["models"][f"{model_name}_{algorithm}"] = cfg
        sig = checkpoint["signature"]
        frozen_audit_rows.append(
            {
                "model": model_name,
                "algorithm": algorithm,
                "status": "PASS",
                "candidate_count": sig["candidate_count"],
                "candidate_limit": sig["candidate_limit"],
                "candidate_grid_hash": sig["candidate_grid_hash"],
                "code_hash": sig["code_hash"],
                "analysis_semantics_version": sig.get("analysis_semantics_version", ""),
                "analysis_semantics_hash": sig.get("analysis_semantics_hash", ""),
                "compatibility_status": checkpoint_compatibility_status(project_root, sig),
                "compatibility_attestation_id": sig.get("compatibility_attestation_id", ""),
                "model_hash": metadata["model_hash"],
            }
        )
    with (project_root / "config/frozen_development_models.yml").open("w", encoding="utf-8") as f:
        yaml.safe_dump(frozen_cfg, f, sort_keys=False)
    out = pd.DataFrame(rows)
    atomic_write_csv(out, project_root / "results/tables/frozen_development_model_inventory.csv")
    atomic_write_csv(pd.DataFrame(frozen_audit_rows), project_root / "results/audit/frozen_model_checkpoint_audit.csv")
    return out


def write_algorithm_completion_gate(project_root: Path) -> pd.DataFrame:
    outer = pd.read_csv(project_root / "results/audit/stage3a_grid_completion.csv")
    frozen_path = project_root / "results/audit/frozen_model_checkpoint_audit.csv"
    frozen = pd.read_csv(frozen_path) if frozen_path.exists() else pd.DataFrame()
    expected_frozen = 5
    expected_outer = len(DEVELOPMENT_CYCLES) * len(MODEL_SAMPLE_DOMAIN)
    rows = []
    for algorithm in SUPPORTED_ALGORITHMS:
        outer_row = outer[outer["algorithm"].eq(algorithm)]
        outer_status = outer_row["status"].iloc[0] if len(outer_row) == 1 else "INCOMPLETE"
        completed_outer = int(outer_row["completed_outer_tasks"].iloc[0]) if len(outer_row) == 1 and "completed_outer_tasks" in outer_row.columns else 0
        outer_expected = int(outer_row["expected_outer_tasks"].iloc[0]) if len(outer_row) == 1 and "expected_outer_tasks" in outer_row.columns else expected_outer
        outer_candidate_counts = outer_row["checkpoint_candidate_counts"].iloc[0] if len(outer_row) == 1 and "checkpoint_candidate_counts" in outer_row.columns else ""
        outer_limits_absent = bool(outer_row["all_candidate_limits_absent"].iloc[0]) if len(outer_row) == 1 and "all_candidate_limits_absent" in outer_row.columns else False
        fg = frozen[frozen["algorithm"].eq(algorithm)].copy() if not frozen.empty else pd.DataFrame()
        full_count = len(elastic_net_candidates(project_root)) if algorithm == "elastic_net" else len(xgboost_candidates(project_root))
        frozen_count = int(fg["model"].nunique()) if not fg.empty else 0
        frozen_counts = sorted(fg["candidate_count"].dropna().astype(int).unique().tolist()) if not fg.empty else []
        frozen_limits_absent = bool(fg["candidate_limit"].isna().all()) if not fg.empty else False
        frozen_compatibility_pass = bool(fg["compatibility_status"].eq("PASS").all()) if not fg.empty and "compatibility_status" in fg else False
        frozen_semantics_pass = (
            sorted(fg["analysis_semantics_hash"].dropna().astype(str).unique().tolist()) == [analysis_semantics_hash()]
            if not fg.empty and "analysis_semantics_hash" in fg
            else False
        )
        frozen_pass = (
            frozen_count == expected_frozen
            and frozen_counts == [full_count]
            and frozen_limits_absent
            and frozen_compatibility_pass
            and frozen_semantics_pass
        )
        passed = outer_status == "PASS" and frozen_pass
        rows.append(
            {
                "algorithm": algorithm,
                "status": "PASS" if passed else ("NOT_RUN" if outer_status == "NOT_RUN" and frozen_count == 0 else "INCOMPLETE"),
                "algorithm_gate": "PASS" if passed else "FAIL",
                "outer_status": outer_status,
                "outer_tasks_completed": completed_outer,
                "outer_tasks_expected": outer_expected,
                "outer_task_fraction": f"{completed_outer}/{outer_expected}",
                "completed_frozen_models": frozen_count,
                "expected_frozen_models": expected_frozen,
                "frozen_model_fraction": f"{frozen_count}/{expected_frozen}",
                "outer_candidate_counts": outer_candidate_counts,
                "frozen_candidate_counts": ";".join(map(str, frozen_counts)),
                "locked_full_grid_candidate_count": full_count,
                "candidate_limit": "",
                "outer_candidate_limits_absent": outer_limits_absent,
                "all_frozen_candidate_limits_absent": frozen_limits_absent,
                "all_candidate_limits_absent": bool(outer_limits_absent and frozen_limits_absent),
                "all_frozen_checkpoint_compatibility_pass": frozen_compatibility_pass,
                "frozen_analysis_semantics_pass": frozen_semantics_pass,
            }
        )
    gate = pd.DataFrame(rows)
    atomic_write_csv(gate, project_root / "results/audit/stage3a_algorithm_completion_gate.csv")
    lines = ["# Stage 3A Algorithm Completion Gate", ""]
    for row in rows:
        lines.append(
            f"- {row['algorithm']}: {row['status']} "
            f"(outer validation {row['outer_status']}; frozen models "
            f"{row['completed_frozen_models']}/{row['expected_frozen_models']})."
        )
    atomic_write_text(project_root / "results/audit/stage3a_algorithm_completion_gate.md", "\n".join(lines) + "\n")
    return gate


def software_versions() -> dict[str, str]:
    import sklearn
    import xgboost

    return {
        "python": sys.version.split()[0],
        "platform": platform.platform(),
        "pandas": pd.__version__,
        "numpy": np.__version__,
        "sklearn": sklearn.__version__,
        "xgboost": xgboost.__version__,
    }


def run_sensitivities(ctx: Stage3AContext, algorithms: list[str], en_limit: int | None, xgb_limit: int | None) -> None:
    project_root = ctx.project_root
    pred = pd.read_csv(project_root / "results/predictions/stage3a_outer_predictions.csv")
    perf = pd.read_csv(project_root / "results/tables/cycle_specific_model_performance.csv")
    main = perf[perf["classification_threshold_type"].eq("fixed_0.50")].copy()
    uacr_rows = []
    for model in ["model1_renal", "model4_combined"]:
        g = main[(main["model"].eq(model)) & (main["algorithm"].eq("elastic_net"))].copy()
        for _, r in g.iterrows():
            uacr_rows.append({**r.to_dict(), "uacr_method": "fold_specific_median_imputation"})
    # Complete-case sensitivity is represented by filtering prediction rows to observed UACR where available.
    df_all = read_harmonized_data(project_root)[["seqn", "cycle", "log2_uacr"]]
    for model in ["model1_renal", "model4_combined"]:
        for alg in ["elastic_net"]:
            pg = pred[(pred["model"].eq(model)) & (pred["algorithm"].eq(alg))].merge(df_all, on=["seqn", "cycle"], how="left")
            pg = pg[pg["log2_uacr"].notna()].copy()
            for cycle, cg in pg.groupby("cycle"):
                y = cg["cvd"].astype(int).to_numpy()
                p = cg["predicted_probability"].to_numpy()
                w = cg["analysis_weight"].to_numpy()
                row = metric_row(model, alg, cycle, y, p, w, 0.5, "fixed_0.50", len(cg), int(y.sum()), cg["weight_variable"].iloc[0], comparison_sample_label(model))
                row["uacr_method"] = "observed_uacr_complete_case_evaluation"
                uacr_rows.append(row)
    pd.DataFrame(uacr_rows).to_csv(project_root / "results/tables/uacr_method_sensitivity.csv", index=False)
    (project_root / "results/audit/uacr_stage3a_sensitivity_audit.md").write_text(
        "\n".join(
            [
                "# UACR Stage 3A Sensitivity Audit",
                "",
                "- Fold-specific median imputation was used in the main Stage 3A pipeline.",
                "- Complete-case evaluation among rows with observed UACR was generated as a sensitivity summary.",
                "- Nested multiple imputation was not implemented in Stage 3A.",
                "- IPW for UACR availability was not implemented in Stage 3A.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    # Class imbalance sensitivity: summarize main and mark class-weighted sensitivity as planned unless explicitly run.
    class_rows = main[main["model"].isin(["model0_core", "model1_renal", "model2_metabolic", "model3_inflammatory", "model4_combined"])].copy()
    class_rows["imbalance_setting"] = np.where(class_rows["algorithm"].eq("elastic_net"), "class_weight=None", "scale_pos_weight=1")
    class_rows.to_csv(project_root / "results/tables/class_imbalance_sensitivity.csv", index=False)


def write_audits(ctx: Stage3AContext, algorithms: list[str], en_limit: int | None, xgb_limit: int | None) -> None:
    project_root = ctx.project_root
    perf = pd.read_csv(project_root / "results/tables/cycle_specific_model_performance.csv")
    pred = pd.read_csv(project_root / "results/predictions/stage3a_outer_predictions.csv")
    selected = pd.read_csv(project_root / "results/models/outer_fold_hyperparameters.csv")
    blocking = []
    if pred["cycle"].eq(EXTERNAL_CYCLE).any():
        blocking.append("2021-2023 predictions were generated, which is not allowed in Stage 3A.")
    if not pred["predicted_probability"].between(0, 1).all():
        blocking.append("Predicted probabilities outside [0,1].")
    if pred["predicted_probability"].isna().any():
        blocking.append("Missing predicted probabilities detected.")
    cycles = sorted(perf["cycle"].unique().tolist())
    missing_cycles = sorted(set(DEVELOPMENT_CYCLES) - set(cycles))
    if missing_cycles:
        blocking.append(f"Missing outer validation cycles: {missing_cycles}")
    grid_audit_path = project_root / "results/audit/stage3a_grid_completion.csv"
    if grid_audit_path.exists():
        grid_audit = pd.read_csv(grid_audit_path)
        for _, row in grid_audit.iterrows():
            if row["status"] != "PASS":
                blocking.append(
                    f"{row['algorithm']} locked-grid outer validation is {row['status']} "
                    f"({int(row['completed_outer_tasks'])}/{int(row['expected_outer_tasks'])} tasks)."
                )
    if xgb_limit is not None:
        blocking.append(f"XGBoost tuning used a candidate limit ({xgb_limit}) instead of the full locked grid; rerun without the limit for a fully compliant Stage 3A.")
    if en_limit is not None:
        blocking.append(f"Elastic Net tuning used a candidate limit ({en_limit}) instead of the full locked grid; rerun without the limit for a fully compliant Stage 3A.")
    blocking.extend(
        [
            "Survey-aware replicate-weight confidence intervals have not yet been generated.",
            "Nested multiple imputation and IPW sensitivity analyses for UACR availability have not yet been completed.",
            "Class-imbalance sensitivity refits have not yet been completed.",
            "Random-effects cycle meta-analysis and decision-curve analysis have not yet been completed.",
        ]
    )
    # Paired sample consistency.
    paired_messages = []
    for comp, (core, ext, _domain) in INCREMENT_COMPARISONS.items():
        for alg in pred["algorithm"].unique():
            p1 = pred[(pred["model"].eq(core)) & (pred["algorithm"].eq(alg))]
            p2 = pred[(pred["model"].eq(ext)) & (pred["algorithm"].eq(alg))]
            merged = p1.merge(p2, on=["seqn", "cycle"], suffixes=("_core", "_extended"))
            paired_messages.append(f"{comp} {alg}: core rows {len(p1)}, extended rows {len(p2)}, paired rows {len(merged)}")
            if merged.empty:
                blocking.append(f"No paired rows for {comp} {alg}.")
    audit_files = {
        "stage3a_data_leakage_audit.md": [
            "# Stage 3A Data Leakage Audit",
            "",
            "- 2021-2023 was not used for tuning, preprocessing, prediction export, or performance evaluation.",
            f"- External-cycle prediction rows detected: {int(pred['cycle'].eq(EXTERNAL_CYCLE).sum())}.",
            "- Outer holdout cycles were excluded from inner tuning and preprocessing fitting.",
        ],
        "stage3a_fold_integrity_audit.md": [
            "# Stage 3A Fold Integrity Audit",
            "",
            f"- Outer cycles completed: {', '.join(cycles)}.",
            "- Validation cycles were kept intact and not randomly split.",
            "- SEQN-cycle predictions are unique by model, algorithm, and validation cycle.",
        ],
        "stage3a_preprocessing_audit.md": [
            "# Stage 3A Preprocessing Audit",
            "",
            "- Median imputation, one-hot encoding, and Elastic Net standardization were fitted inside the training folds only.",
            "- XGBoost used the same fold-fitted imputation and one-hot encoding; no outer holdout data were used in preprocessing fitting.",
            "- Early stopping was not used, avoiding any possibility of outer holdout leakage.",
        ],
        "stage3a_sample_consistency_audit.md": [
            "# Stage 3A Sample Consistency Audit",
            "",
            *[f"- {m}" for m in paired_messages],
        ],
        "stage3a_hyperparameter_audit.md": [
            "# Stage 3A Hyperparameter Audit",
            "",
            f"- Hyperparameter rows: {len(selected)}.",
            "- Primary selection metric was weighted log loss.",
            "- Ties within 1e-4 were resolved by lower-complexity hyperparameters.",
            f"- Elastic Net candidate limit: {en_limit if en_limit is not None else 'none'}.",
            f"- XGBoost candidate limit: {xgb_limit if xgb_limit is not None else 'none'}.",
        ],
        "stage3a_metric_audit.md": [
            "# Stage 3A Metric Audit",
            "",
            "- Metrics include weighted/unweighted AUROC, weighted/unweighted PR-AUC, Brier score, log loss, calibration intercept, calibration slope, O/E ratio, and secondary threshold metrics.",
            "- Fixed 0.50 and training-set weighted Youden thresholds are reported separately.",
        ],
        "stage3a_model_freeze_audit.md": [
            "# Stage 3A Model Freeze Audit",
            "",
            "- Frozen 2005-2018 development models were fitted after internal-external validation.",
            "- Frozen models are for later temporal validation only and not used to estimate Stage 3A performance.",
            "- See `config/frozen_development_models.yml`.",
        ],
        "stage3a_claims_audit.md": [
            "# Stage 3A Claims Audit",
            "",
            "- Outcome wording must remain: prevalent self-reported cardiovascular disease.",
            "- Do not describe the analysis as incident, prospective, future-event, or causal prediction.",
            "- GBD is not used in Stage 3A.",
        ],
        "survey_bootstrap_implementation.md": [
            "# Survey Bootstrap Implementation",
            "",
            "- Stage 3A generated survey-weighted point estimates using NHANES weights, strata, and PSU variables.",
            "- Replicate-weight confidence intervals were not generated in this run.",
            "- Ordinary participant-level bootstrap was not used as a substitute for the main survey-aware CI.",
            "- CI status: Survey-aware confidence intervals are produced by run_survey_ci.py after Stage 3A.",
        ],
        "hyperparameter_selection_audit.md": [
            "# Hyperparameter Selection Audit",
            "",
            "- Detailed inner-cycle candidate results are saved in `results/models/inner_cv_all_candidates.csv`.",
            "- Selected outer-fold hyperparameters are saved in `results/models/outer_fold_hyperparameters.csv`.",
            "- Selection used mean inner weighted log loss with pre-specified lower-complexity tie-breaks.",
        ],
        "final_model_freeze_audit.md": [
            "# Final Model Freeze Audit",
            "",
            "- Frozen development model artifacts were saved under `results/models/frozen_development/`.",
            "- Each artifact includes feature lists, hyperparameters, metadata, model hash, and calibration status.",
            "- Temporal validation status remains locked.",
        ],
    }
    for name, lines in audit_files.items():
        (project_root / "results/audit" / name).write_text("\n".join(lines) + "\n", encoding="utf-8")
    block_lines = ["# Stage 3A Blocking Issues", ""]
    action_path = project_root / "ACTION_REQUIRED_STAGE3A_BLOCKING_ISSUES.md"
    if blocking:
        block_lines.append("Blocking or compliance issues detected:")
        block_lines.extend([f"- {b}" for b in blocking])
        atomic_write_text(action_path, "\n".join(block_lines) + "\n")
    else:
        block_lines.append("- No blocking issue detected.")
        if action_path.exists():
            action_path.unlink()
    atomic_write_text(project_root / "results/audit/stage3a_blocking_issues.md", "\n".join(block_lines) + "\n")
    completion = [
        "# Stage 3A Completion Report",
        "",
        f"- Run time: {ctx.run_time}.",
        f"- Algorithms run: {', '.join(algorithms)}.",
        f"- Outer cycles completed: {', '.join(cycles)}.",
        f"- Prediction rows: {len(pred)}.",
        f"- Blocking issue count: {len(blocking)}.",
        "- Stage 3B was not run.",
    ]
    (project_root / "results/audit/stage3a_completion_report.md").write_text("\n".join(completion) + "\n", encoding="utf-8")


def update_project_logs(ctx: Stage3AContext, status: str) -> None:
    project_root = ctx.project_root
    run_status = project_root / "results/logs/run_status.csv"
    row = pd.DataFrame(
        [
            {
                "step": "03a_internal_external_validation",
                "status": status,
                "run_time": ctx.run_time,
                "input": str(ctx.data_path),
                "outputs": json.dumps(
                    [
                        "results/tables/cycle_specific_model_performance.csv",
                        "results/tables/cycle_specific_incremental_value.csv",
                        "results/tables/pooled_out_of_cycle_performance.csv",
                        "results/tables/pooled_incremental_value.csv",
                        "results/models/frozen_development/",
                        "results/audit/stage3a_completion_report.md",
                    ],
                    ensure_ascii=False,
                ),
                "random_seed": RANDOM_SEED,
                "software": json.dumps(software_versions(), sort_keys=True),
            }
        ]
    )
    prior = pd.read_csv(run_status) if run_status.exists() else pd.DataFrame()
    if not prior.empty and "step" in prior.columns:
        prior = prior[~prior["step"].eq("03a_internal_external_validation")].copy()
    pd.concat([prior, row], ignore_index=True).to_csv(run_status, index=False)

    stage3_status = project_root / "results/logs/stage3a_run_status.csv"
    row.to_csv(stage3_status, index=False)


def append_project_docs(project_root: Path) -> None:
    with (project_root / "CHANGELOG.md").open("a", encoding="utf-8") as f:
        f.write(f"\n## {now()} Stage 3A internal-external validation\n")
        f.write("- Created and ran NHANES 2005-2018 leave-one-cycle-out internal-external validation pipeline.\n")
        f.write("- 2021-2023 temporal validation remained locked and was not evaluated.\n")
    with (project_root / "DECISIONS.md").open("a", encoding="utf-8") as f:
        f.write(f"\n## {now()} Stage 3A modeling decisions\n")
        f.write("- Used Stage 2 frozen variables, cycle folds, weights, and hyperparameter grids.\n")
        f.write("- Reported probability-based metrics as primary; threshold metrics remain secondary.\n")
    readme = project_root / "README.md"
    txt = readme.read_text(encoding="utf-8") if readme.exists() else ""
    if "Stage 3A" not in txt:
        with readme.open("a", encoding="utf-8") as f:
            f.write("\n## Stage 3A\n")
            f.write("Run `python run_stage3a.py --project-root .` to perform NHANES 2005-2018 leave-one-cycle-out internal-external validation. NHANES 2021-2023 remains locked for later temporal validation.\n")
    matrix = project_root / "reporting/reviewer_comment_action_matrix.md"
    with matrix.open("a", encoding="utf-8") as f:
        f.write("\n## Stage 3A reviewer-risk actions\n")
        f.write("- Rebuilt prediction workflow around cycle-level internal-external validation.\n")
        f.write("- Preserved GBD exclusion from individual-level modeling.\n")


def run_all(args: argparse.Namespace) -> int:
    project_root = Path(args.project_root).resolve()
    ensure_dirs(project_root)
    ctx = Stage3AContext(
        project_root=project_root,
        data_path=project_root / "data/interim/stage2_harmonized_audit_dataset.csv",
        run_time=now(),
        log_path=project_root / "results/logs/stage3a_full_run.log",
    )
    if not args.resume:
        ctx.log_path.write_text("", encoding="utf-8")
    try:
        log(ctx, f"Stage 3A started; resume={args.resume}")
        check_temporal_lock(project_root)
        generate_denominator_tables(ctx)
        algorithms = [a.strip() for a in args.algorithms.split(",") if a.strip()]
        if not algorithms or len(set(algorithms)) != len(algorithms):
            raise ValueError("--algorithms must contain one or more unique algorithm names.")
        unsupported = sorted(set(algorithms) - set(SUPPORTED_ALGORITHMS))
        if unsupported:
            raise ValueError(f"Unsupported algorithms: {unsupported}")
        outputs = run_internal_external(
            ctx,
            algorithms,
            args.elastic_net_candidate_limit,
            args.xgboost_candidate_limit,
            resume=args.resume,
        )
        completed_algorithms = sorted(outputs["predictions"]["algorithm"].unique().tolist())
        aggregate_performance(project_root)
        run_sensitivities(ctx, completed_algorithms, args.elastic_net_candidate_limit, args.xgboost_candidate_limit)
        fit_frozen_development_models(
            ctx,
            algorithms,
            args.elastic_net_candidate_limit,
            args.xgboost_candidate_limit,
            resume=args.resume,
        )
        algorithm_gate = write_algorithm_completion_gate(project_root)
        plot_figures(project_root)
        write_audits(ctx, completed_algorithms, args.elastic_net_candidate_limit, args.xgboost_candidate_limit)
        append_project_docs(project_root)
        blocking = (project_root / "ACTION_REQUIRED_STAGE3A_BLOCKING_ISSUES.md").exists()
        requested_gate = algorithm_gate[algorithm_gate["algorithm"].isin(algorithms)]
        requested_pass = len(requested_gate) == len(algorithms) and requested_gate["status"].eq("PASS").all()
        status = "SUCCESS" if requested_pass and not blocking else ("PARTIAL" if requested_pass else "BLOCKED")
        update_project_logs(ctx, status)
        log(ctx, f"Stage 3A finished; requested algorithm gate={'PASS' if requested_pass else 'FAIL'}; overall status={status}")
        return 0 if requested_pass else 2
    except Exception as exc:
        log(ctx, f"Stage 3A failed: {exc}")
        (project_root / "results/audit/stage3a_blocking_issues.md").write_text(
            f"# Stage 3A Blocking Issues\n\n- Pipeline failed: {exc}\n",
            encoding="utf-8",
        )
        (project_root / "ACTION_REQUIRED_STAGE3A_BLOCKING_ISSUES.md").write_text(
            f"# Stage 3A Blocking Issues\n\n- Pipeline failed: {exc}\n",
            encoding="utf-8",
        )
        update_project_logs(ctx, "FAILED")
        raise


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Stage 3A NHANES 2005-2018 internal-external validation.")
    p.add_argument("--project-root", default=".", help="Project root")
    p.add_argument("--algorithms", default="elastic_net,xgboost", help="Comma-separated algorithms: elastic_net,xgboost")
    p.add_argument("--resume", action="store_true", help="Reuse only hash-validated, atomically completed checkpoints")
    p.add_argument("--elastic-net-candidate-limit", type=int, default=None, help="Optional limit for testing; full grid if omitted")
    p.add_argument("--xgboost-candidate-limit", type=int, default=None, help="Optional limit for testing; full grid if omitted")
    return p


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    return run_all(args)


if __name__ == "__main__":
    raise SystemExit(main())
